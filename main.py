from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import sys
from PyQt5.Qt import QApplication, QClipboard
from PyQt5 import QtCore, QtWidgets
from PyQt5.QtWidgets import QMainWindow, QWidget, QPlainTextEdit, QCheckBox, QLabel, QPushButton, QFileDialog
from PyQt5.QtCore import QSize

mydoc = Document()

pre_words = """Acknowledging
Acting
Affirming
Alarmed by
Alarmed
Anxious
Appreciating
Approving
Aware of
Bearing in mind
Believing
Cognizant
Concerned
Confident
Conscious
Considering
Contemplating
Convinced
Declaring
Deeply concerned
Deeply conscious
Deeply convinced
Deeply disturbed
Deeply regretting
Deploring
Desiring
Determined
Emphasizing
Encouraged
Expecting
Expressing appreciation
Noting with approval
Expressing concern also
Expressing concern
Expressing its appreciation
Expressing its satisfaction
Expressing satisfaction
Firmly convinced
Fulfilling
Fully alarmed
Fully aware
Fully believing
Further deploring
Further recalling
Guided by
Having adopted
Having considered
Having considered further
Having devoted attention
Having examined
Having heard
Having received
Having reviewed
Having studied
Having adopted
Having approved
Having considered
Having decided
Keeping in mind
Mindful
Noting
Noting further
Noting with deep concern
Noting with regret
Noting with satisfaction
Observing
Reaffirming
Reaffirming also
Realizing
Recalling
Recalling also
Recognizing
Recognizing also
Recognizing with satisfaction
Referring
Regretting
Reiterating
Reiterating its call for
Reminding
Seeking
Seized
Stressing
Taking into account
Taking into consideration
Taking note
Taking note also
Taking note further
Underlining
Viewing with appreciation
Viewing with apprehension
Welcoming
Welcoming also"""

op_words = """Accepts
Acknowledges
Adopts
Advises
Affirms
Also calls for
Also recommends
Also strongly condemns
Also urges
Appeals
Appreciates
Approves
Authorizes
Calls
Calls for
Calls upon
Commends
Concurs
Condemns
Confirms
Congratulates
Considers
Decides
Declares
Declares accordingly
Demands
Deplores
Designates
Directs
Draws the attention
Emphasizes
Encourages
Endorses
Expresses its appreciation
Expresses its hope
Expresses its regret
Further invites
Further proclaims
Further recommends
Further reminds
Further requests
Further resolves
Has resolved
Instructs
Introduces
Invites
Notes
Notes with satisfaction
Proclaims
Reaffirms
Recalls
Recognizes
Recommends
Regrets
Reiterates
Reminds
Renews its appeal
Repeats
Requests
Requires
Solemnly affirms
Stresses
Strongly advises
Strongly condemns
Strongly encourages
Suggests
Supports
Takes note of
Transmits
Trusts
Underlines
Underscores
Urges
Welcomes"""

def split_pre(k):
    final = [k[0:k.index(" ")], k[k.index(" ")::]]
    for j in pre_words.splitlines():
        kk = j.strip().lower()
        
        if k[0:len(kk)].lower() == kk:
            
            final[0] = j.strip()
            final[1] = k[len(kk)::]

    return final

def split_op(k):
    
    final = [k[0:k.index(" ")], k[k.index(" ")::]]
    for j in op_words.splitlines():
        kk = j.strip().lower()
        
        if k[0:len(kk)].lower() == kk:
            
            final[0] = j.strip()
            final[1] = k[len(kk)::]

    return final

def write(text, size, is_bold, indent):

    global mydoc 

    para = mydoc.add_paragraph("")
    paragraph_format = para.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    
    if indent != None:
        paragraph_format.left_indent = Inches(indent)
    k = para.add_run(text)
    k.font.name = 'Times New Roman'
    k.font.size = Pt(size)
    k.font.bold = is_bold

def write_preamb(kk, b, it, u):
    final =split_pre(kk)

    if final[0][0].upper() != final[0][0]:
        final[0] = final[0][0].upper() + final[0][1::]
    
    if final[1][len(final[1]) - 1] == "," or final[1][len(final[1]) - 1] == ";" or final[1][len(final[1]) - 1] == ".":
        final[1] = final[1][0:len(final[1]) - 1] + "," 
    else:
        final[1] += ","
        
    global mydoc
    para = mydoc.add_paragraph("")
    paragraph_format = para.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    k = para.add_run(final[0])
    k.font.name = 'Times New Roman'
    k.font.size = Pt(12)

    k.font.bold = b
    if u == True:
        k.font.underline = True
    k.font.italic = it

    k2 = para.add_run(final[1])
    k2.font.name = 'Times New Roman'
    k2.font.size = Pt(12)

def write_op(kk, number, b, it, u, end):
    final = split_op(kk)
    
    
    if final[0][0].upper() != final[0][0]:
        final[0] = final[0][0].upper() + final[0][1::]
    
    if final[1][len(final[1]) - 1] == "," or final[1][len(final[1]) - 1] == ";" or final[1][len(final[1]) - 1] == "." or final[1][len(final[1]) - 1] == "-":
        final[1] = final[1][0:len(final[1]) - 1] + end 
    else:
        final[1] += end
        
    global mydoc
    para = mydoc.add_paragraph("")
    paragraph_format = para.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    k0 = para.add_run(str(number) + ") ")
    k0.font.name = 'Times New Roman'
    k0.font.size = Pt(12)
    k = para.add_run(final[0])
    k.font.name = 'Times New Roman'
    k.font.size = Pt(12)

    k.font.bold = b
    if u == True:
        k.font.underline = True
    k.font.italic = it

    k2 = para.add_run(final[1])
    k2.font.name = 'Times New Roman'
    k2.font.size = Pt(12)

def is_space(k):
    for i in k:
        if i == " ":
            return True 
    return False

letters = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z']
def write_sub(k, number, end):
    x = k

    if k[0].upper == k[0]:
        x = k[0].lower() + k[1::]

    if k[len(k) - 1] == ";" or k[len(k) - 1] == "," or k[len(k) - 1] == "."  or k[len(k) - 1] == "-":
        x = x[0:len(x) - 1] + end
    else:
        x += end

    x = letters[number - 1] + ") " + x 

    write(x, 12, False, 0.5)

roman_numbers = ['i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x', 'xi', 'xii', 'xiii', 'xiv']

def write_sub_sub(k, number, end):
    x = k

    if k[0].upper == k[0]:
        x = k[0].lower() + k[1::]

    if k[len(k) - 1] == ";" or k[len(k) - 1] == "," or k[len(k) - 1] == "." or k[len(k) - 1] == "-":
        x = x[0:len(x) - 1] + end
    else:
        x += end

    x = roman_numbers[number - 1] + ") " + x 

    write(x, 12, False, 1)

def remove_full_stops(k):
    final = ""

    for i in k:
        if i != ".":
            final += i
    return final

def update(title, committee, topic, sponsors, signatories, preambs, operatives, b1, i1, u1, b2, i2, u2, name):
    global mydoc
    mydoc = Document()
    write(title, 16, True, 0)
    
    committee = remove_full_stops(committee)
    topic = remove_full_stops(topic)
    sponsors = remove_full_stops(sponsors)

    signatories = remove_full_stops(signatories)
    preambs = remove_full_stops(preambs)
    operatives = remove_full_stops(operatives)

    
    write("", 12, False, 0)
    write("Committee: " + committee, 12, False, 0)
    write("Topic: " + topic, 12, False, 0)
    write("Sponsors: " + sponsors, 12, False, 0)
    write("Signatories: " + signatories, 12, False, 0)
    write("", 12, False, 0)
    write("The " + committee.strip() + ",", 12, False, 0)
    for i in preambs.splitlines():
        if is_space(i) == False:
            continue
        else:
            if i != "":
                write("", 12, False, 0)
                write_preamb(i, b1, i1, u1)
            
    clause_number = 0
    sub_clause_number = 0
    sub_sub_clause_number = 0
    yy = operatives.splitlines()
    y = []
    for i in yy:
        if i != "" and len(i) > 2 and is_space(i) == True:
            y.append(i)

    if len(y) == 0:
        
        return
    for i in range(0, len(y) - 1):
        if y[i] != "":
            if y[i][0] == "*" and y[i][1] == "*":
                sub_sub_clause_number += 1
                if y[i + 1][0] != "*":
                    write_sub_sub(y[i][2::].strip(), sub_sub_clause_number, ';')
                else:
                    write_sub_sub(y[i][2::].strip(), sub_sub_clause_number, ',')
            elif y[i][0] == "*":
                sub_clause_number += 1
                if y[i + 1][0] != "*":
                    write_sub(y[i][1::].strip(), sub_clause_number, ';')
                else:
                    if y[i + 1][1] == "*":
                        write_sub(y[i][1::].strip(), sub_clause_number, '-')
                    else:
                        write_sub(y[i][1::].strip(), sub_clause_number, ',')
            else:
                clause_number += 1
                sub_clause_number = 0
                sub_sub_clause_number = 0
                write("", 12, False, 0)
                if y[i + 1][0] == "*":
                    write_op(y[i].strip(), clause_number, b2, i2, u2, '-')
                else:
                    write_op(y[i].strip(), clause_number, b2, i2, u2, ';')    
            
    y_final = y[len(y) - 1]
    if y_final[0] == "*" and y_final[1] == "*":
        write_sub_sub(y_final[2::].strip(), sub_sub_clause_number + 1, '.')
    elif y_final[0] == "*":
        write_sub(y_final[1::].strip(), sub_clause_number + 1, '.')
    else:
        write("", 12, False, 0)
        write_op(y_final.strip(), clause_number + 1, b2, i2, u2, '.')
    mydoc.save(name)

def save():
    name, _ = QFileDialog.getSaveFileName(w,'Save File','', 'docx files (*.docx)')
    if len(name) > 0:
        update(title.toPlainText(), committee.toPlainText(), topic.toPlainText(), sponsors.toPlainText(), signatories.toPlainText(), pre.toPlainText(), op.toPlainText(), preamb_bold.isChecked(), preamb_italic.isChecked(), preamb_underlined.isChecked(), op_bold.isChecked(), op_italic.isChecked(), op_underlined.isChecked(), name)


        

appctxt = QApplication([])
w = QWidget()  
w.setWindowTitle("Reso Writer")
w.resize(1460, 940)
 

title = QPlainTextEdit(w)
title.setPlaceholderText("Enter title of resolution")
title.move(10, 10)
title.resize(945, 39)

sponsors = QPlainTextEdit(w)
sponsors.setPlaceholderText("Enter sponsor nations (separated by comma and space)")
sponsors.move(10, 160)
sponsors.resize(945, 40)
    
signatories = QPlainTextEdit(w)
signatories.setPlaceholderText("Enter signatory nations (separated by comma and space)")
signatories.move(10, 210)
signatories.resize(945, 40)

committee = QPlainTextEdit(w)
committee.setPlaceholderText("Enter committee name")
committee.move(10, 60)
committee.resize(945, 40)

topic = QPlainTextEdit(w)
topic.setPlaceholderText("Enter topic/agenda")
topic.move(10, 110)
topic.resize(945, 40)

l1 = QLabel("Preambulatory clause formatting:", w)
l1.resize(300, 20)
l1.move(10, 260)
        
preamb_bold = QCheckBox("Bold", w)
preamb_bold.move(220, 256)

preamb_italic = QCheckBox("Italic", w)
preamb_italic.move(280, 256)

preamb_underlined = QCheckBox("Underlined", w)
preamb_underlined.move(340, 256)

l2 = QLabel("Operative clause formatting:", w)
l2.resize(300, 20)
l2.move(10, 290)
        
op_bold = QCheckBox("Bold", w)
op_bold.move(220, 286)

op_italic = QCheckBox("Italic", w)
op_italic.move(280, 286)

op_underlined = QCheckBox("Underlined", w)
op_underlined.move(340, 286)

pre = QPlainTextEdit(w)
pre.setPlaceholderText("Enter preambulatory clauses")
pre.move(10,330)
pre.resize(1440, 156)

op = QPlainTextEdit(w)
op.setPlaceholderText("Enter operative clauses")
op.move(10,496)
op.resize(1440, 395)

submit = QPushButton("Export", w)
submit.move(10, 900)
submit.clicked.connect(save)

    
w.show()
sys.exit(appctxt.exec_())